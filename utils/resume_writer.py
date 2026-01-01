




from docx import Document
from docx.shared import Pt

def strip_hashes(text: str) -> str:
    return text.replace("#", "")

def strip_stars(text: str) -> str:
    return text.replace("*", "")

def extract_between_markers(text: str) -> str:
    start = text.find("---")
    end = text.rfind("---")

    if start == -1 or end == -1 or start == end:
        return text.strip()  # fallback if markers missing

    return text[start + 3 : end].strip()


def add_markdown_bold_paragraph(doc, text):
   
    paragraph = doc.add_paragraph()
    parts = text.split("**")

    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        run.font.size = Pt(10.5)

        if i % 2 == 1:
            run.bold = True

def save_enhanced_resume(text: str, filename="enhanced_resume.docx"):
    doc = Document()

    for line in strip_hashes(text).split("\n"):
        if not line.strip():
            doc.add_paragraph()
            continue

      
        add_markdown_bold_paragraph(doc, line)

        
        paragraph = doc.paragraphs[-1]
        if line.isupper() or line.endswith(":"):
            for run in paragraph.runs:
                run.font.size = Pt(12)
        else:
            for run in paragraph.runs:
                run.font.size = Pt(10.5)

    

    doc.save(filename)
